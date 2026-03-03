"""Parser for DOCX files using python-docx.

Supports: .docx files (Microsoft Word format).
Extracts: Full text, tables as Markdown, heading-based sections, metadata.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import structlog

from intake.ingest.base import ParsedContent, ParseError, validate_file_readable

logger = structlog.get_logger()


class DocxParser:
    """Parser for Microsoft Word (.docx) documents.

    Supports:
    - .docx files with text, headings, and tables
    - Heading-based section extraction
    - Table conversion to Markdown

    Extracts:
    - Full text from paragraphs
    - Tables as Markdown-formatted text
    - Sections based on heading styles
    - Document metadata (author, title)
    """

    def can_parse(self, source: str) -> bool:
        """Check if this source is a parseable DOCX file."""
        path = Path(source)
        return path.exists() and path.is_file() and path.suffix.lower() == ".docx"

    def parse(self, source: str) -> ParsedContent:
        """Parse a DOCX file into normalized content.

        Args:
            source: Path to the DOCX file.

        Returns:
            ParsedContent with extracted text, sections, and metadata.

        Raises:
            ParseError: If the DOCX cannot be read or is empty.
        """
        path = validate_file_readable(source)

        try:
            from docx import Document
        except ImportError as e:
            raise ParseError(
                source=source,
                reason="python-docx is not installed",
                suggestion="Install it with: pip install python-docx",
            ) from e

        try:
            doc = Document(str(path))
        except Exception as e:
            raise ParseError(
                source=source,
                reason=f"Could not open DOCX: {e}",
                suggestion="Verify the file is a valid .docx document.",
            ) from e

        text_parts: list[str] = []
        sections = self._extract_sections(doc)
        tables_md = self._extract_tables(doc)
        metadata = self._extract_metadata(doc)

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        full_text = "\n\n".join(text_parts)
        if tables_md:
            full_text += "\n\n" + tables_md

        if not full_text.strip():
            raise ParseError(
                source=source,
                reason="DOCX document contains no extractable text",
                suggestion="Verify the document has text content (not just images).",
            )

        metadata["source_type"] = "docx"

        logger.info(
            "docx_parsed",
            source=source,
            paragraphs=len(text_parts),
            sections=len(sections),
            tables=tables_md.count("|---") if tables_md else 0,
        )

        return ParsedContent(
            text=full_text,
            format="docx",
            source=source,
            metadata=metadata,
            sections=sections,
        )

    def _extract_sections(self, doc: Any) -> list[dict[str, str]]:
        """Extract sections based on heading paragraph styles."""
        sections: list[dict[str, str]] = []
        current_heading: str | None = None
        current_level: str = "0"
        current_content: list[str] = []

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""

            if style_name.startswith("Heading"):
                if current_heading is not None:
                    sections.append(
                        {
                            "title": current_heading,
                            "level": current_level,
                            "content": "\n".join(current_content).strip(),
                        }
                    )

                current_heading = paragraph.text.strip()
                level_part = style_name.replace("Heading", "").strip()
                current_level = level_part if level_part.isdigit() else "1"
                current_content = []
            elif paragraph.text.strip():
                current_content.append(paragraph.text)

        if current_heading is not None:
            sections.append(
                {
                    "title": current_heading,
                    "level": current_level,
                    "content": "\n".join(current_content).strip(),
                }
            )

        return sections

    def _extract_tables(self, doc: Any) -> str:
        """Extract all tables from the document as Markdown."""
        parts: list[str] = []
        for table in doc.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(cells)

            if not rows:
                continue

            md = self._rows_to_markdown(rows)
            if md:
                parts.append(md)

        return "\n\n".join(parts)

    def _rows_to_markdown(self, rows: list[list[str]]) -> str:
        """Convert table rows to Markdown format."""
        if not rows:
            return ""

        header = rows[0]
        lines: list[str] = []
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")

        for row in rows[1:]:
            while len(row) < len(header):
                row.append("")
            lines.append("| " + " | ".join(row[: len(header)]) + " |")

        return "\n".join(lines)

    def _extract_metadata(self, doc: Any) -> dict[str, str]:
        """Extract document metadata (author, title, etc.)."""
        metadata: dict[str, str] = {}
        props = doc.core_properties

        if props.author:
            metadata["author"] = props.author
        if props.title:
            metadata["title"] = props.title
        if props.subject:
            metadata["subject"] = props.subject
        if props.created:
            metadata["created"] = str(props.created)

        return metadata
